from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


class DocumentExtractionError(ValueError):
    pass


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int | None
    text: str


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_document(filename: str, content: bytes) -> list[ExtractedPage]:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractionError(
            "Unsupported file type. Upload PDF, DOCX, TXT, or Markdown files."
        )
    if not content:
        raise DocumentExtractionError("The uploaded file is empty.")

    if extension == ".pdf":
        pages = _extract_pdf(content)
    elif extension == ".docx":
        pages = _extract_docx(content)
    else:
        pages = _extract_plain_text(content)

    cleaned = [
        ExtractedPage(page.page_number, _clean_text(page.text))
        for page in pages
        if _clean_text(page.text)
    ]
    if not cleaned:
        raise DocumentExtractionError(
            "No readable text was found. Scanned PDFs require an OCR service before ingestion."
        )
    return cleaned


def _extract_pdf(content: bytes) -> list[ExtractedPage]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:  # pragma: no cover - library-specific
                raise DocumentExtractionError("Password-protected PDFs are not supported.") from exc
        return [
            ExtractedPage(index, page.extract_text() or "")
            for index, page in enumerate(reader.pages, start=1)
        ]
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError(f"Could not read PDF: {exc}") from exc


def _extract_docx(content: bytes) -> list[ExtractedPage]:
    try:
        from docx import Document as DocxDocument

        document = DocxDocument(BytesIO(content))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style:
                blocks.append(f"\n{text}\n")
            else:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return [ExtractedPage(None, "\n\n".join(blocks))]
    except Exception as exc:
        raise DocumentExtractionError(f"Could not read DOCX: {exc}") from exc


def _extract_plain_text(content: bytes) -> list[ExtractedPage]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = content.decode("latin-1")
    logical_pages = text.split("\f")
    if len(logical_pages) == 1:
        return [ExtractedPage(None, text)]
    return [ExtractedPage(index, page) for index, page in enumerate(logical_pages, start=1)]


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
